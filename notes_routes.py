import re
import io
from flask import Blueprint, jsonify, session, send_file, request
from database import get_db_connection, log_user_activity
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# BLUEPRINT SETUP
# ============================================================================

notes_bp = Blueprint('notes', __name__)


# ============================================================================
# ROUTES
# ============================================================================

@notes_bp.route('/api/documents/<int:document_id>/notes', methods=['GET'])
def get_document_notes(document_id):
    """
    Return the raw markdown notes text for a specific document.
    Used by the Notes panel to display lecture content.
    """
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    user_id = session['user_id']
    conn = get_db_connection()

    row = conn.execute(
        'SELECT id, original_filename, markdown_text FROM documents WHERE id = ? AND user_id = ?',
        (document_id, user_id)
    ).fetchone()

    conn.close()

    if row is None:
        return jsonify({'success': False, 'message': 'Document not found'}), 404

    markdown = row['markdown_text'] or ''

    if not markdown:
        return jsonify({
            'success': False,
            'message': 'No notes available for this document. It may have been uploaded before this feature was added.',
        }), 404

    log_user_activity(session['user_id'], session.get('username'), 'notes_view',
                      detail=f"Viewed notes: {row['original_filename']}",
                      ip_address=request.remote_addr)
    return jsonify({
        'success': True,
        'document_id': document_id,
        'filename': row['original_filename'],
        'markdown': markdown,
    })


@notes_bp.route('/api/documents/<int:document_id>/notes/download', methods=['GET'])
def download_document_notes(document_id):
    """
    Convert the stored markdown notes to a .docx file and send it
    as a download. Uses python-docx to build the document.
    """
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    user_id = session['user_id']
    conn = get_db_connection()

    row = conn.execute(
        'SELECT original_filename, markdown_text FROM documents WHERE id = ? AND user_id = ?',
        (document_id, user_id)
    ).fetchone()

    conn.close()

    if row is None:
        return jsonify({'success': False, 'message': 'Document not found'}), 404

    markdown = row['markdown_text'] or ''
    if not markdown:
        return jsonify({'success': False, 'message': 'No notes available for this document.'}), 404

    # ── Build the .docx ──
    doc = DocxDocument()

    # Title heading
    title = doc.add_heading(row['original_filename'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacer

    for line in markdown.split('\n'):
        stripped = line.rstrip()

        # Headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)

        # Unordered list
        elif re.match(r'^[\*\-\+]\s', stripped):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, stripped[2:].strip())

        # Ordered list
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            _add_inline(p, text)

        # Horizontal rule
        elif re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', stripped):
            doc.add_paragraph('─' * 60)

        # Blockquote
        elif stripped.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            _add_inline(p, stripped[2:].strip())

        # Blank line → spacer
        elif stripped == '':
            doc.add_paragraph()

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)

    # ── Stream to response ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r'[^\w\-.]', '_', row['original_filename'].rsplit('.', 1)[0])
    download_name = f"{safe_name}_notes.docx"

    log_user_activity(session['user_id'], session.get('username'), 'notes_download',
                      detail=f"Downloaded notes for: {row['original_filename']}",
                      ip_address=request.remote_addr)
    return send_file(
        buf,
        as_attachment=True,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ============================================================================
# HELPERS
# ============================================================================

def _add_inline(paragraph, text):
    """
    Parse inline markdown (bold, italic, inline code) and add
    styled runs to a python-docx paragraph.
    """
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)')
    last = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith('**'):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif full.startswith('*'):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif full.startswith('`'):
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        last = m.end()
    # Remaining plain text
    if last < len(text):
        paragraph.add_run(text[last:])